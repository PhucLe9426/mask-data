"""Extract text from supported uploads without sending the original file elsewhere."""

from io import BytesIO
from itertools import islice
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader
import xlrd


SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".xml",
    ".yaml",
    ".yml",
    ".log",
    ".pdf",
    ".docx",
    ".xlsx",
    ".xls",
}
MAX_EXCEL_COLUMNS = 512


class FileExtractionError(ValueError):
    """Raised when an upload cannot be safely converted to text."""


def safe_filename(filename: str | None) -> str:
    name = Path(filename or "file").name.strip()
    return name[:255] or "file"


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "cp1258"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise FileExtractionError("Không đọc được bảng mã của tệp văn bản. Hãy lưu tệp dưới dạng UTF-8.")


def _extract_pdf(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content))
        if reader.is_encrypted:
            raise FileExtractionError("PDF đang được bảo vệ bằng mật khẩu.")
        return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)
    except FileExtractionError:
        raise
    except Exception as exc:
        raise FileExtractionError("Không thể đọc tệp PDF hoặc PDF đã bị hỏng.") from exc


def _extract_docx(content: bytes, max_uncompressed_bytes: int) -> str:
    try:
        with ZipFile(BytesIO(content)) as archive:
            unpacked_size = sum(item.file_size for item in archive.infolist())
            if unpacked_size > max_uncompressed_bytes:
                raise FileExtractionError("Tệp DOCX quá lớn sau khi giải nén.")

        document = Document(BytesIO(content))
        parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            # python-docx trả cùng một XML cell nhiều lần cho các ô được merge.
            # Chỉ đọc mỗi ô thật một lần để biểu mẫu Word không bị nhân bản nội dung.
            seen_cells: set[int] = set()
            for row in table.rows:
                values: list[str] = []
                for cell in row.cells:
                    cell_id = id(cell._tc)
                    if cell_id in seen_cells:
                        continue
                    seen_cells.add(cell_id)
                    value = cell.text.strip()
                    if value:
                        values.append(value)
                if any(values):
                    parts.append(" | ".join(values))
        return "\n".join(parts)
    except FileExtractionError:
        raise
    except Exception as exc:
        raise FileExtractionError("Không thể đọc tệp DOCX hoặc DOCX đã bị hỏng.") from exc


def _append_excel_row(parts: list[str], values, max_chars: int, current_size: int) -> int:
    rendered = [
        str(value).strip() if value is not None else ""
        for value in islice(values, MAX_EXCEL_COLUMNS)
    ]
    while rendered and not rendered[-1]:
        rendered.pop()
    if not rendered:
        return current_size
    line = "\t".join(rendered)
    next_size = current_size + len(line) + 1
    if next_size > max_chars:
        raise FileExtractionError(
            f"Nội dung sau khi trích xuất vượt quá {max_chars:,} ký tự. Hãy chia tệp thành phần nhỏ hơn."
        )
    parts.append(line)
    return next_size


def _extract_xlsx(content: bytes, max_uncompressed_bytes: int, max_chars: int) -> str:
    try:
        with ZipFile(BytesIO(content)) as archive:
            unpacked_size = sum(item.file_size for item in archive.infolist())
            if unpacked_size > max_uncompressed_bytes:
                raise FileExtractionError("Tệp Excel quá lớn sau khi giải nén.")

        workbook = load_workbook(
            BytesIO(content), read_only=True, data_only=False, keep_links=False
        )
        parts: list[str] = []
        current_size = 0
        data_rows = 0
        try:
            for sheet in workbook.worksheets:
                # Một số file do Excel/phần mềm kế toán tạo khai báo vùng dùng tới
                # hàng 1.048.576 hoặc cột XFD dù chỉ có vài ô. Trong read-only mode,
                # bỏ dimension không chính xác để parser chỉ đi qua cell thật trong XML.
                sheet.reset_dimensions()
                parts.append(f"--- Sheet: {sheet.title} ---")
                for row in sheet.iter_rows(values_only=True):
                    next_size = _append_excel_row(parts, row, max_chars, current_size)
                    if next_size != current_size:
                        data_rows += 1
                        current_size = next_size
        finally:
            workbook.close()
        return "\n".join(parts) if data_rows else ""
    except FileExtractionError:
        raise
    except Exception as exc:
        raise FileExtractionError("Không thể đọc tệp XLSX hoặc tệp đã bị hỏng/mã hóa.") from exc


def _extract_xls(content: bytes, max_chars: int) -> str:
    try:
        workbook = xlrd.open_workbook(file_contents=content, on_demand=True)
        parts: list[str] = []
        current_size = 0
        data_rows = 0
        try:
            for sheet in workbook.sheets():
                parts.append(f"--- Sheet: {sheet.name} ---")
                for row_index in range(sheet.nrows):
                    next_size = _append_excel_row(
                        parts, sheet.row_values(row_index), max_chars, current_size
                    )
                    if next_size != current_size:
                        data_rows += 1
                        current_size = next_size
        finally:
            workbook.release_resources()
        return "\n".join(parts) if data_rows else ""
    except FileExtractionError:
        raise
    except Exception as exc:
        raise FileExtractionError("Không thể đọc tệp XLS hoặc tệp đã bị hỏng/mã hóa.") from exc


def extract_text(
    filename: str | None,
    content: bytes,
    *,
    max_bytes: int,
    max_chars: int,
) -> tuple[str, str]:
    """Return a sanitized filename and locally extracted text."""
    name = safe_filename(filename)
    extension = Path(name).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise FileExtractionError(f"Định dạng tệp chưa được hỗ trợ. Các định dạng hợp lệ: {supported}")
    if not content:
        raise FileExtractionError("Tệp đang trống.")
    if len(content) > max_bytes:
        limit_mb = max_bytes / (1024 * 1024)
        raise FileExtractionError(f"Tệp vượt quá giới hạn {limit_mb:g} MB.")

    if extension == ".pdf":
        text = _extract_pdf(content)
    elif extension == ".docx":
        text = _extract_docx(content, max_uncompressed_bytes=max_bytes * 5)
    elif extension == ".xlsx":
        text = _extract_xlsx(content, max_uncompressed_bytes=max_bytes * 10, max_chars=max_chars)
    elif extension == ".xls":
        text = _extract_xls(content, max_chars=max_chars)
    else:
        text = _decode_text(content)

    text = text.replace("\x00", "").strip()
    if not text:
        if extension == ".pdf":
            raise FileExtractionError("PDF không có lớp văn bản. Hiện ứng dụng chưa hỗ trợ OCR PDF scan.")
        raise FileExtractionError("Không tìm thấy nội dung văn bản trong tệp.")
    if len(text) > max_chars:
        raise FileExtractionError(
            f"Nội dung sau khi trích xuất vượt quá {max_chars:,} ký tự. Hãy chia tệp thành phần nhỏ hơn."
        )
    return name, text
