"""Create downloadable conversation exports without writing sensitive data to disk."""

from __future__ import annotations

from functools import lru_cache
from html import escape
from io import BytesIO
from pathlib import Path
import re
import unicodedata

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


EXPORT_TYPES = {
    "docx": (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    ),
    "xlsx": (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "xlsx",
    ),
    "pdf": ("application/pdf", "pdf"),
}
INVALID_XML_CHARS = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")


def _clean_text(value: str) -> str:
    return INVALID_XML_CHARS.sub("", str(value))


def safe_filename(title: str, extension: str) -> str:
    normalized = unicodedata.normalize("NFKD", title).encode("ascii", "ignore").decode()
    stem = re.sub(r"[^a-zA-Z0-9]+", "-", normalized).strip("-").lower()
    return f"{stem[:80] or 'tong-hop-ai'}.{extension}"


def _source_values(source: dict) -> tuple[str, str]:
    return _clean_text(source.get("name") or "Tài liệu"), _clean_text(source.get("excerpt") or "")


def build_docx(title: str, content: str, sources: list[dict]) -> bytes:
    title = _clean_text(title)
    content = _clean_text(content)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    heading = document.add_heading(title, level=1)
    heading.runs[0].font.color.rgb = RGBColor(24, 151, 125)

    for line in content.splitlines() or [content]:
        document.add_paragraph(line)

    if sources:
        document.add_heading("Nguồn tham khảo", level=2)
        for source in sources:
            name, excerpt_text = _source_values(source)
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(name).bold = True
            if excerpt_text:
                paragraph.add_run(f" — {excerpt_text}")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _excel_text(value: str) -> str:
    """Force untrusted LLM output to remain text instead of an Excel formula."""
    value = _clean_text(value)
    if value.startswith(("=", "+", "-", "@")):
        return "'" + value
    return value


def _content_chunks(content: str, limit: int = 30_000) -> list[str]:
    chunks: list[str] = []
    for line in content.splitlines() or [content]:
        if not line:
            chunks.append("")
            continue
        chunks.extend(line[index:index + limit] for index in range(0, len(line), limit))
    return chunks


def build_xlsx(title: str, content: str, sources: list[dict]) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Tổng hợp"
    sheet.sheet_view.showGridLines = False
    sheet.merge_cells("A1:D1")
    sheet["A1"] = _excel_text(title)
    sheet["A1"].font = Font(size=16, bold=True, color="FFFFFF")
    sheet["A1"].fill = PatternFill("solid", fgColor="18977D")
    sheet["A1"].alignment = Alignment(vertical="center")
    sheet.row_dimensions[1].height = 32

    sheet.merge_cells("A3:D3")
    sheet["A3"] = "Nội dung tổng hợp"
    sheet["A3"].font = Font(size=12, bold=True, color="18212F")
    sheet["A3"].fill = PatternFill("solid", fgColor="E8F3EF")

    row = 4
    for chunk in _content_chunks(content):
        sheet.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        cell = sheet.cell(row=row, column=1, value=_excel_text(chunk))
        cell.alignment = Alignment(wrap_text=True, vertical="top")
        sheet.row_dimensions[row].height = min(120, max(22, 16 * (1 + len(chunk) // 95)))
        row += 1

    for column, width in enumerate((34, 24, 24, 24), start=1):
        sheet.column_dimensions[get_column_letter(column)].width = width
    sheet.freeze_panes = "A4"
    sheet.print_area = f"A1:D{max(4, row - 1)}"
    sheet.sheet_properties.pageSetUpPr.fitToPage = True
    sheet.page_setup.fitToWidth = 1
    sheet.page_setup.fitToHeight = 0

    if sources:
        source_sheet = workbook.create_sheet("Nguồn")
        source_sheet.sheet_view.showGridLines = False
        source_sheet.append(["Tên tài liệu", "Đoạn trích"])
        for cell in source_sheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="18977D")
        for source in sources:
            name, excerpt_text = _source_values(source)
            source_sheet.append([_excel_text(name), _excel_text(excerpt_text[:30_000])])
        source_sheet.column_dimensions["A"].width = 36
        source_sheet.column_dimensions["B"].width = 100
        for cells in source_sheet.iter_rows(min_row=2):
            for cell in cells:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
        source_sheet.freeze_panes = "A2"

    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


@lru_cache(maxsize=1)
def _register_pdf_fonts() -> tuple[str, str]:
    regular_candidates = (
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    )
    bold_candidates = (
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf"),
    )
    regular = next((path for path in regular_candidates if path.exists()), None)
    bold = next((path for path in bold_candidates if path.exists()), None)
    if regular:
        pdfmetrics.registerFont(TTFont("ExportText", str(regular)))
    if bold:
        pdfmetrics.registerFont(TTFont("ExportTextBold", str(bold)))
    return ("ExportText" if regular else "Helvetica", "ExportTextBold" if bold else "Helvetica-Bold")


def build_pdf(title: str, content: str, sources: list[dict]) -> bytes:
    title = _clean_text(title)
    content = _clean_text(content)
    regular_font, bold_font = _register_pdf_fonts()
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=title,
        author="Data Masking",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ExportTitle",
        parent=styles["Title"],
        fontName=bold_font,
        fontSize=20,
        leading=25,
        textColor=colors.HexColor("#18977D"),
        alignment=TA_LEFT,
        spaceAfter=12,
    )
    body_style = ParagraphStyle(
        "ExportBody",
        parent=styles["BodyText"],
        fontName=regular_font,
        fontSize=10.5,
        leading=16,
        textColor=colors.HexColor("#18212F"),
        spaceAfter=7,
    )
    source_heading = ParagraphStyle(
        "ExportSourceHeading",
        parent=body_style,
        fontName=bold_font,
        fontSize=13,
        spaceBefore=12,
        spaceAfter=8,
    )

    story = [Paragraph(escape(title), title_style)]
    for line in content.splitlines() or [content]:
        story.append(Paragraph(escape(line) or "&nbsp;", body_style))
    if sources:
        story.extend((Spacer(1, 5 * mm), Paragraph("Nguồn tham khảo", source_heading)))
        for source in sources:
            name, excerpt_text = _source_values(source)
            source_line = f"<b>{escape(name)}</b>"
            if excerpt_text:
                source_line += f" — {escape(excerpt_text)}"
            story.append(Paragraph(f"• {source_line}", body_style))
    document.build(story)
    return output.getvalue()


def build_export(file_format: str, title: str, content: str, sources: list[dict]) -> tuple[bytes, str, str]:
    if file_format not in EXPORT_TYPES:
        raise ValueError("Định dạng xuất file không được hỗ trợ")
    builders = {"docx": build_docx, "xlsx": build_xlsx, "pdf": build_pdf}
    mime_type, extension = EXPORT_TYPES[file_format]
    data = builders[file_format](title.strip() or "Tổng hợp AI", content, sources)
    return data, mime_type, safe_filename(title, extension)
